# -*- coding: utf-8 -*-
"""
Build 数字乡村平台需求规格说明书.docx
- 章节：产品背景/目标/建设范围/用户/业务流程/功能需求/数据需求/配置/非功能/技术架构/运维部署/商业模式/实施说明/附录(术语表)
- 功能需求按：县级/乡级/村级/大屏/村民移动/村干部/通用
- 每个功能：功能说明 + 功能截图(全图) + 功能字段及描述(表) + 是否有业务流转
"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = r"D:/dev/pro_digitalvillages"
OUT_DIR = os.path.join(ROOT, "build")
DATA = os.path.join(OUT_DIR, "data.json")
SHOT_DIR = os.path.join(OUT_DIR, "screenshots")
DIAGRAM_DIR = os.path.join(OUT_DIR, "diagrams", "png")
DOCX = os.path.join(ROOT, "docs", "需求规格说明书.docx")

with open(DATA, encoding="utf-8") as fh:
    _raw = json.load(fh)
# normalize keys to be relative to html/ so they match CATALOG rels
DATA_JSON = {k.replace("html/", ""): v for k, v in _raw.items()}

# ---------- CJK font helper ----------
def set_cjk(run, font="宋体"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)

def para(doc, text, size=10.5, bold=False, font="宋体", align=None, color=None, space_after=4, space_before=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    set_cjk(r, font)
    return p

def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13.5, 3: 12}
    fonts = {1: "黑体", 2: "黑体", 3: "黑体"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level==1 else 6)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.paragraph_format.page_break_before = False
    r = p.add_run(text)
    r.font.size = Pt(sizes[level]); r.bold = True
    set_cjk(r, fonts[level])
    # outline level for TOC
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level-1)); pPr.append(ol)
    return p

def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（在 Word 中右键“更新域”即可刷新目录）"
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    r._r.append(fldBegin); r._r.append(instr); r._r.append(fldSep); r._r.append(t); r._r.append(fldEnd)

def sublabel(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); set_cjk(r, "黑体")
    return p

# ---------- images ----------
def add_image(doc, rel, width_px=600):
    name = rel.replace("html/", "").replace("/", "_").replace(".html", ".png")
    path = os.path.join(SHOT_DIR, name)
    if not os.path.exists(path):
        para(doc, "（截图缺失：%s）" % rel, size=9, color=RGBColor(0x99,0x33,0x33))
        return
    try:
        Image.open(path)
    except Exception:
        para(doc, "（截图无法读取：%s）" % rel, size=9); return
    width_in = width_px / 96.0
    doc.add_picture(path, width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(2)
    dd = DATA_JSON.get(rel, {})
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("图：%s" % (dd.get("main_title") or dd.get("title_tag") or rel))
    cr.font.size = Pt(8.5); set_cjk(cr, "宋体"); cr.font.color.rgb = RGBColor(0x66,0x66,0x66)

def add_image_list(doc, rels):
    for rel in rels:
        add_image(doc, rel, width_px=600 if DATA_JSON[rel]["device"] in ("admin","dashboard","shared") else 250)

def insert_diagram(doc, png_name, caption):
    path = os.path.join(DIAGRAM_DIR, png_name)
    if not os.path.exists(path):
        para(doc, "（图表缺失：%s）" % png_name, size=9, color=RGBColor(0x99,0x33,0x33))
        return
    doc.add_picture(path, width=Cm(15.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(2)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.size = Pt(9); set_cjk(cr, "宋体"); cr.font.color.rgb = RGBColor(0x33,0x33,0x33)

# ---------- field table ----------
def field_table(doc, rel):
    d = DATA_JSON.get(rel)
    if d is None:
        para(doc, "（原型页面未找到：%s，请核对路径）" % rel, size=9, color=RGBColor(0x99,0x33,0x33))
        return
    fields = d.get("fields", [])
    tables = d.get("tables", [])
    if fields:
        rows = [["字段名称", "类型", "必填", "只读", "说明 / 选项"]]
        for f in fields:
            opt = ""
            if f["options"]:
                opt = "选项：" + "、".join(f["options"][:12]) + ("…" if len(f["options"])>12 else "")
            desc = f["placeholder"] or opt or "信息录入/展示项"
            if opt and f["placeholder"]:
                desc = f["placeholder"] + "；" + opt
            rows.append([f["label"], f["type"], "是" if f["required"] else "否",
                         "是" if f["readonly"] else "否", desc])
        make_table(doc, rows, [Cm(3.0), Cm(1.6), Cm(1.2), Cm(1.2), Cm(8.5)])
        return
    if tables:
        t = tables[0]
        rows = [["序号", "列表/表格字段", "说明"]]
        for i, col in enumerate(t["header"], 1):
            rows.append([str(i), col, "数据列展示项"])
        make_table(doc, rows, [Cm(1.5), Cm(6.0), Cm(7.0)])
        if len(tables) > 1:
            para(doc, "（本页面另含 %d 个数据表，字段以截图为凭）" % (len(tables)-1), size=9, color=RGBColor(0x66,0x66,0x66))
        return
    # fallback: sections
    secs = d.get("sections", [])
    if secs:
        rows = [["序号", "模块 / 指标项"]]
        for i, s in enumerate(secs, 1):
            rows.append([str(i), s])
        make_table(doc, rows, [Cm(1.5), Cm(13.0)])
        para(doc, "（本页面以大屏数据卡片/图表为主，核心指标见截图）", size=9, color=RGBColor(0x66,0x66,0x66))
        return
    para(doc, "（本页面为门户/导航页，无独立录入字段，详见截图）", size=9, color=RGBColor(0x66,0x66,0x66))

def make_table(doc, rows, widths_cm):
    ncol = len(rows[0])
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
            r = p.add_run(str(val))
            r.font.size = Pt(9 if r_i > 0 else 9.5)
            if r_i == 0:
                r.bold = True; set_cjk(r, "黑体")
                shade_cell(cells[c_i], "DCE6F1")
            else:
                set_cjk(r, "宋体")
            # width
            if c_i < len(widths_cm):
                cells[c_i].width = widths_cm[c_i]
    # set column widths on table grid
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        # widths already set per cell; also set grid
    except Exception:
        pass

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# ---------- function catalog ----------
# section key: county / town / village / dashboard / villager / cadre / common
CATALOG = {
 "county": [
  ("县级工作台（管理驾驶舱）", ["admin/index.html","admin/dashboard.html"],
   "县级运营总览门户，聚合待办事项、关键数据指标、预警信息与快捷操作入口，是县级管理员日常工作的统一 landing 页。", False, ""),
  ("通知公告管理", ["admin/notice/notice-list.html"],
   "县级统一发布面向全县/乡镇/村的政务通知、政策公告与公示信息，支持列表查看、检索与发布状态管理。", True, "业务流转：起草 → 审核 → 发布 → 各端（村民/干部）接收查看。"),
  ("村务（事务）管理", ["admin/affairs/affairs-list.html"],
   "对村级事务上报、审批与归档的统一管理，支撑县对村务处理进度的监督与督办。", True, "业务流转：村级提交事务 → 乡/县审核 → 归档跟踪。"),
  ("群众诉求（信访）管理", ["admin/appeal/appeal-list.html"],
   "集中处理群众通过移动端提交的诉求/信访工单，支持分派、办理与办结回访。", True, "业务流转：群众提交诉求 → 县级分派 → 责任单位办理 → 办结反馈。"),
  ("村民议事（话题）管理", ["admin/discuss/discuss-list.html"],
   "对村民移动端发起的议事话题进行后台审核、置顶与归档，保障议事内容合规有序。", True, "业务流转：村民发起话题 → 后台审核 → 公开讨论 → 形成决议。"),
  ("活动/事件管理", ["admin/event/event-list.html"],
   "县级统筹发布文体、宣教、惠民等活动与突发事件登记，便于统一调度与统计。", True, "业务流转：活动创建 → 发布报名 → 执行 → 复盘。"),
  ("工作日志管理", ["admin/journal/journal-list.html"],
   "汇总各村镇干部工作日志记录，支持检索、点评与考核留痕。", False, ""),
  ("报表管理", ["admin/report/report-list.html"],
   "县级各类业务报表的集中管理与下发，支持按周期生成与导出。", False, ""),
  ("便民服务事项管理", ["admin/service/service-guide.html","admin/service/service-apply.html"],
   "维护县级便民服务事项目录与办事指南，并管理群众线上申请的受理流转。", True, "业务流转：维护指南 → 群众申请 → 受理 → 办理 → 办结。"),
  ("数据统计与分析", ["admin/stats/report.html"],
   "面向县级决策者的多维度数据统计看板，覆盖人口、产业、治理、服务等领域。", False, ""),
  ("任务管理", ["admin/task/task-list.html"],
   "县级向乡镇、村下发重点工作任务的创建、分派、跟踪与考核。", True, "业务流转：创建任务 → 分派 → 执行反馈 → 验收考核。"),
  ("组织/机构管理（组织树）", ["admin/user/org-tree.html"],
   "维护县—乡—村三级组织架构与隶属关系，是权限与数据隔离的基础。", False, ""),
  ("角色与权限管理", ["admin/user/role-permission.html"],
   "定义县级及各下属机构的角色与菜单/数据权限，落实最小权限原则。", False, ""),
  ("村干部花名册", ["admin/user/cadre-list.html"],
   "县级统一管理各乡镇、村干部基本信息与任职状态，支撑干部考核与联络。", False, ""),
  ("村民档案管理", ["admin/user/villager-list.html"],
   "县级汇聚各村村民基础档案，支撑精准服务与统计分析。", False, ""),
  ("县域开通审核（乡镇/村入驻）", ["admin/city/town-apply.html","admin/city/village-apply.html","admin/city/open-audit.html","admin/city/operator-list.html"],
   "新乡镇、新村接入平台时的入驻申请、资料审核与开通运营全流程管理。", True, "业务流转：乡镇/村提交入驻申请 → 县级审核资料 → 开通账号与配置 → 正式运营。"),
  ("数据大屏配置（数据源/内容/规则）", ["admin/datascreen/data-source.html","admin/datascreen/content-config.html","admin/datascreen/display-rule.html"],
   "县级对数据大屏的数据接入、展示内容与轮播规则的集中配置，实现“一屏统览”。", True, "业务流转：配置数据源 → 配置展示内容 → 配置播放规则 → 大屏生效。"),
  ("系统参数配置", ["admin/system/config.html"],
   "平台级通用参数（如积分规则开关、消息模板、地图密钥等）的配置。", False, ""),
  ("数据字典管理", ["admin/system/dict.html"],
   "维护全平台统一的枚举字典（如证件类型、事项分类），保证数据一致。", False, ""),
  ("操作日志审计", ["admin/system/log.html"],
   "记录管理员关键操作日志，支撑安全审计与责任追溯。", False, ""),
 ],
 "town": [
  ("乡镇概况管理", ["admin/town/my-town/info.html"],
   "维护本乡镇基本概况、简介与展示信息，作为乡镇级门户与上级数据来源。", False, ""),
  ("乡镇产业数据管理", ["admin/town/data/industry-data.html"],
   "采集与维护本乡镇产业（种植、养殖、加工、旅游等）基础数据与产值。", False, ""),
  ("乡镇党建组织管理", ["admin/town/party/org-manage.html"],
   "管理乡镇级党组织架构、党员信息与组织生活记录。", False, ""),
 ],
 "village": [
  ("村级工作台（概览）", ["admin/village/my-village/overview.html"],
   "村级管理员日常工作的统一入口，展示本村关键指标、待办与快捷功能。", False, ""),
  ("村民（农户）档案管理", ["admin/village/my-village/farmer-manage.html"],
   "维护本村农户/村民详细档案，含家庭、承包、补贴等多维信息，是精准治理基础。", False, ""),
  ("村联系人管理", ["admin/village/my-village/contacts.html"],
   "维护村两委、网格员、党员中心户等联系人及联络方式。", False, ""),
  ("村工作队/专班管理", ["admin/village/my-village/team.html"],
   "管理驻村工作队、帮扶专班成员及其职责分工。", False, ""),
  ("贫困户/帮扶对象管理", ["admin/village/my-village/poverty-household.html"],
   "建档立卡脱贫户、监测户与帮扶措施的动态管理。", True, "业务流转：识别登记 → 制定帮扶措施 → 跟踪帮扶 → 效果评估。"),
  ("村入驻/成员审核", ["admin/village/my-village/register-audit.html"],
   "对本村新入驻成员、认领关系等进行审核确认。", True, "业务流转：申请 → 村级审核 → 通过/驳回。"),
  ("三务公开（村务公开）管理", ["admin/village/sun/affairs-public.html","admin/village/sun/four-open-detail.html"],
   "维护党务、村务、财务“三务公开”内容发布与详情，落实阳光村务。", True, "业务流转：起草公开内容 → 审核 → 发布公示 → 村民查看。"),
  ("党建信息发布", ["admin/village/party/info-publish.html"],
   "村级党组织发布党建动态、通知与学习资料。", True, "业务流转：编辑 → 发布 → 党员/群众查看。"),
  ("村产业（种养殖）管理", ["admin/village/data/industry-manage.html"],
   "维护本村特色产业项目、基地与经营主体信息。", False, ""),
  ("土地资源/地块管理", ["admin/village/data/land.html"],
   "管理村内耕地、林地、宅基地等地块矢量与属性信息。", False, ""),
  ("房屋/宅基地管理", ["admin/village/data/house.html"],
   "维护村民房屋、宅基地台账与户型信息。", False, ""),
  ("村设备/物联网管理", ["admin/village/device/device-manage.html"],
   "管理村内监控、环境传感、广播等物联网设备台账与状态。", False, ""),
  ("认养农业管理", ["admin/village/adopt/adopt-manage.html"],
   "发布并管理“我在乡下有块田”等农业认养项目，连接城市消费者与农户。", True, "业务流转：发布认养项目 → 消费者认养 → 种养管理 → 产出配送。"),
  ("乡村旅游管理", ["admin/village/tourism/tourism-manage.html"],
   "维护村级旅游资源、景点、民宿与农旅活动信息。", False, ""),
  ("积分商城管理", ["admin/village/points-shop/points-manage.html","admin/village/points-shop/goods-detail.html","admin/village/points-shop/order-detail.html"],
   "村级运营积分商城的商品的上下架、库存与兑换订单处理。", True, "业务流转：商品上架 → 村民兑换 → 订单处理 → 发货/核销。"),
  ("表单应用管理", ["admin/village/form/form-manage.html"],
   "通过低代码表单能力快速配置村级业务采集表单（如摸底表、报名表）。", False, ""),
  ("千村千面设计器", ["admin/village/designer/home-designer.html"],
   "为不同村定制首页布局、配色与模块组合的“千村千面”可视化设计工具。", False, ""),
  ("治理-矛盾调解", ["admin/village/governance/mediation.html"],
   "登记、流转与归档村级矛盾纠纷调解案件。", True, "业务流转：受理登记 → 调解 → 达成协议/上报 → 结案。"),
  ("治理-微网格管理", ["admin/village/governance/micro-grid.html"],
   "划分村级微网格、配置网格员并管理网格事件。", True, "业务流转：网格划分 → 事件上报 → 网格员处置 → 上报闭环。"),
  ("治理-流动人口管理", ["admin/village/governance/migrant.html"],
   "登记与管理本村外来流动人口信息，支撑治安与服务。", False, ""),
  ("村务审核管理", ["admin/village/audit/audit-manage.html"],
   "对村级提交的关键业务（如公开、认养、积分）进行复核审批。", True, "业务流转：提交 → 审核 → 通过/驳回 → 执行。"),
 ],
 "dashboard": [
  ("县级数据大屏（总览）", ["dashboard/index.html"],
   "县级指挥调度大屏，一屏统览全县乡村治理、产业、民生核心指标与地理分布。", False, ""),
  ("乡镇数据大屏", ["dashboard/town/overview.html"],
   "乡镇级数据大屏，展示本乡镇概况、产业、治理与服务的实时态势。", False, ""),
  ("村级-文化大屏", ["dashboard/village/culture.html"],
   "村级文化宣传大屏，展示村史、乡风文明、文化活动等内容。", False, ""),
  ("村级-数据中心大屏", ["dashboard/village/data-center.html"],
   "村级数据汇聚大屏，集中呈现人口、土地、产业等基础数据。", False, ""),
  ("村级-经济大屏", ["dashboard/village/economy.html"],
   "展示村级集体经济、农户收入、产业产值等经济指标。", False, ""),
  ("村级-治理大屏", ["dashboard/village/governance.html"],
   "展示村级网格治理、事件处置、平安建设等治理成效。", False, ""),
  ("村级-党建大屏", ["dashboard/village/party.html"],
   "展示村级党组织、党员、组织生活与党建活动。", False, ""),
  ("村级-服务大屏", ["dashboard/village/service.html"],
   "展示村级便民服务办件量、满意度与热点服务事项。", False, ""),
 ],
 "villager": [
  ("村民首页", ["villager/index.html"],
   "村民移动端首页，聚合通知、服务入口、积分与村务动态，是村民最高频入口。", False, ""),
  ("村务/三务查看", ["villager/affairs/affairs-list.html","villager/affairs/affairs-detail.html"],
   "村民查看本村党务、村务、财务公开内容，保障知情权与监督权。", False, ""),
  ("我的设备/物联网", ["villager/device/device-list.html","villager/device/device-detail.html"],
   "村民查看绑定到本人/家庭的物联网设备（如监测、广播）及状态。", False, ""),
  ("议事话题", ["villager/discuss/topic-list.html","villager/discuss/topic-detail.html","villager/discuss/topic-new.html"],
   "村民发起、参与村级议事话题讨论，促进民主决策。", True, "业务流转：发起话题 → 邻里讨论 → 后台审核 → 形成决议。"),
  ("积分首页与明细", ["villager/points/points-home.html","villager/points/exchange-log.html"],
   "展示村民当前积分余额、获取记录与消耗明细，激励参与治理。", False, ""),
  ("积分商城", ["villager/points/points-mall.html","villager/points/points-exchange.html"],
   "村民使用积分兑换实物/服务，支撑乡村治理“积分制”。", True, "业务流转：挑选商品 → 积分兑换 → 生成订单 → 线下核销/发货。"),
  ("问题随手拍上报", ["villager/report/report-home.html","villager/report/report-form.html","villager/report/report-detail.html","villager/report/my-reports.html","villager/report/report-success.html"],
   "村民对村内环境问题、设施损坏等一键拍照上报，跟踪处理进度。", True, "业务流转：拍照上报 → 系统分派 → 干部处置 → 结果反馈 → 评价。"),
  ("便民服务", ["villager/service/service-home.html","villager/service/guide-list.html","villager/service/guide-detail.html","villager/service/apply-form.html","villager/service/apply-success.html","villager/service/my-records.html"],
   "村民在线查阅办事指南并在线提交各类便民申请，实现“掌上办”。", True, "业务流转：查指南 → 在线申请 → 受理审核 → 办理 → 结果回执。"),
 ],
 "cadre": [
  ("干部首页", ["cadre/index.html"],
   "村干部移动端工作台，聚合待办任务、巡查、上报与积分审核入口。", False, ""),
  ("诉求处理", ["cadre/appeal/appeal-list.html","cadre/appeal/appeal-detail.html"],
   "村干部接收并办理群众诉求工单，跟踪处置与办结。", True, "业务流转：接收工单 → 核实办理 → 办结反馈。"),
  ("数据概览", ["cadre/data/data-overview.html"],
   "村干部查看本村关键业务数据概览，辅助日常管理决策。", False, ""),
  ("设备管理", ["cadre/device/device-list.html","cadre/device/device-detail.html"],
   "村干部查看与巡检本村物联网设备状态与告警。", False, ""),
  ("工作日志", ["cadre/journal/journal-list.html","cadre/journal/journal-detail.html","cadre/journal/journal-write.html"],
   "村干部撰写、提交与查看工作日志，沉淀履职记录。", True, "业务流转：撰写日志 → 提交 → 上级点评/归档。"),
  ("巡查巡检", ["cadre/patrol/patrol-plan.html","cadre/patrol/patrol-checkin.html","cadre/patrol/patrol-records.html","cadre/patrol/patrol-detail.html","cadre/patrol/patrol-stats.html"],
   "村干部制定巡查计划、现场打卡、记录事件并统计巡查成效。", True, "业务流转：制定计划 → 到点打卡 → 记录事件 → 汇总统计。"),
  ("积分审核", ["cadre/points/audit-list.html","cadre/points/audit-detail.html","cadre/points/audit-history.html"],
   "村干部审核村民积分获取/兑换申请，把关积分发放。", True, "业务流转：提交积分申请 → 干部审核 → 发放/驳回 → 留痕。"),
  ("问题上报处理", ["cadre/report/report-home.html","cadre/report/report-form.html","cadre/report/report-success.html"],
   "村干部侧的问题上报与处置入口（与村民上报协同）。", True, "业务流转：上报/接收 → 处置 → 反馈。"),
  ("任务执行", ["cadre/task/task-list.html","cadre/task/task-detail.html","cadre/task/task-execute.html"],
   "村干部接收上级下发任务、执行反馈与提交办结。", True, "业务流转：接收任务 → 执行 → 反馈进度 → 提交办结。"),
 ],
 "common": [
  ("统一身份认证与登录", ["admin/login.html","villager/login.html","cadre/login.html","villager/auth.html"],
   "各端统一的账号登录、注册与身份鉴权体系，基于角色与组织隔离数据。", False, ""),
  ("个人中心（资料管理）", ["villager/profile/profile.html","villager/profile/profile-edit.html","cadre/profile/profile.html","cadre/profile/profile-edit.html"],
   "用户维护个人资料、账号与安全设置，跨端一致的身份视图。", False, ""),
  ("消息通知中心", ["villager/notice/notice-list.html","villager/notice/notice-detail.html","cadre/notice/notice-sent.html","cadre/notice/notice-send.html"],
   "统一消息通知的接收与发送，连接县级发布与村民/干部触达。", True, "业务流转：后台/干部发布 → 系统推送 → 用户查看。"),
  ("帮助与反馈", ["villager/profile/feedback.html","villager/profile/about.html"],
   "提供使用帮助、关于信息与用户意见反馈入口。", False, ""),
  ("前端组件规范库", ["shared/components-demo.html"],
   "统一 UI 组件与样式规范演示，保障四端体验一致性。", False, ""),
 ],
}

SECTION_TITLES = {
 "county": "6.1 县级平台功能",
 "town": "6.2 乡级平台功能",
 "village": "6.3 村级平台功能",
 "dashboard": "6.4 大屏端功能",
 "villager": "6.5 村民移动端功能",
 "cadre": "6.6 村干部功能",
 "common": "6.7 通用功能",
}

# =========================================================
# Build document
# =========================================================
doc = Document()
# default font
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr(); rf = OxmlElement('w:rFonts')
rf.set(qn('w:eastAsia'),'宋体'); rpr.append(rf)
# page size A4
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0)

def H1(t): heading(doc, t, 1)
def H2(t): heading(doc, t, 2)
def H3(t): heading(doc, t, 3)

# ---------- Cover ----------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run("数字乡村平台"); r.font.size = Pt(30); r.bold=True; set_cjk(r,"黑体")
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("需求规格说明书（SRS）"); r2.font.size = Pt(22); r2.bold=True; set_cjk(r2,"黑体")
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(30)
r3 = p3.add_run("—— 覆盖县 / 乡 / 村三级治理与村民、干部、大屏多端 ——"); r3.font.size = Pt(12); set_cjk(r3,"宋体")
for _ in range(2): doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("文档版本：V1.0\n编制日期：2026-08-12\n编制角色：产品经理（产品通）\n文档状态：评审稿"); rm.font.size = Pt(11); set_cjk(rm,"宋体")
doc.add_page_break()

# ---------- TOC ----------
H1("目录")
add_toc(doc)
doc.add_page_break()

# ---------- 1 产品背景 ----------
H1("1. 产品背景")
H2("1.1 政策背景")
para(doc, "数字乡村是乡村振兴的战略方向。国家《“十四五”数字经济发展规划》《数字乡村发展战略纲要》明确提出以数字技术驱动乡村治理现代化、产业数字化与公共服务均等化。县级统建、乡村共用的数字化平台，已成为落实“治理下沉、服务上门”的关键基础设施。")
H2("1.2 行业现状与痛点")
para(doc, "当前县域乡村治理普遍存在：信息孤岛严重，县-乡-村数据不贯通；村务不透明，群众监督缺位；基层干部事务繁杂、留痕困难；便民服务“多头跑、反复跑”；产业与文旅资源缺乏线上载体。本平台以“三级协同、多端联动”破解上述痛点。")
H2("1.3 数字化转型需求")
para(doc, "建设统一数字乡村平台，实现：组织在线（三级组织树）、村务在线（三务公开）、治理在线（网格+随手拍）、服务在线（掌上办）、产业在线（认养+文旅）、展示在线（数据大屏），并以积分制激活群众参与。")

# ---------- 2 产品目标 ----------
H1("2. 产品目标")
H2("2.1 总体目标")
para(doc, "打造“一平台、三级用、多端触达”的数字乡村操作系统：县级统管、乡级枢纽、村级落地，覆盖 PC 管理后台、村干部/村民移动端、数据大屏四端，实现治理增效、产业升级、服务便民。")
H2("2.2 分级目标")
para(doc, "县级：统建平台、数据贯通、监督考核、决策大屏；乡级：承上启下、产业统筹、组织管理；村级：落地执行、三务公开、网格治理、积分运营；群众/干部：移动办事、移动履职。")
H2("2.3 量化指标（建议）")
make_table(doc, [["维度","目标（上线首年）"],
 ["村级入驻率","≥ 95% 行政村上线"],
 ["便民事项掌上办占比","≥ 80%"],
 ["群众诉求办结时效","平均 ≤ 3 个工作日"],
 ["大屏覆盖","县/乡/重点村 100%"],
 ["日活（村民端）","覆盖村民 ≥ 60%"]], [Cm(6), Cm(8)])
H2("2.4 非目标（Non-goals）")
para(doc, "本期不实现：财政/医保等强业务系统对接（仅做服务入口）、AI 智能决策模型、跨县跨省数据交换、硬件生产制造。")

# ---------- 3 建设范围 ----------
H1("3. 建设范围")
H2("3.1 系统总体架构边界")
para(doc, "平台以“前端四端 + 业务中台 + 数据底座”为总体边界，本说明书聚焦四端原型所覆盖的业务功能与配套配置、运维、部署要求。")
insert_diagram(doc, "business_architecture.png", "图 3-1 数字乡村平台业务架构图")
H2("3.2 三级管理边界")
make_table(doc, [["层级","职责边界","对应原型模块"],
 ["县级","统建、审核、监督、考核、大屏配置","admin（县级模块）/city/datascreen"],
 ["乡级","产业统筹、组织管理、承上启下","admin/town"],
 ["村级","落地执行、三务公开、网格治理、积分","admin/village"]], [Cm(2), Cm(8), Cm(4)])
H2("3.3 四端定位")
para(doc, "PC 管理后台（admin）：县/乡/村三级运营中枢；村干部移动端（cadre）：移动履职；村民移动端（villager）：掌上办事；数据大屏端（dashboard）：指挥展示。")
H2("3.4 功能覆盖范围总表")
para(doc, "详见第 6 章功能需求。全量原型页面 123 个，已按 7 个分组归类，截图与字段见各功能小节。")
H2("3.5 与现有原型对应关系")
para(doc, "本说明书功能需求章节所列每一个功能均对应 html/ 目录下真实原型页面，并附全图截图与字段清单，确保需求与实现一致、可追溯。")

# ---------- 4 产品用户 ----------
H1("4. 产品用户")
H2("4.1 用户角色定义")
make_table(doc, [["角色","说明","主要使用端"],
 ["县级管理员","统管全县平台与数据","PC 后台"],
 ["乡镇管理员","本乡镇运营与产业统筹","PC 后台"],
 ["村级管理员","本村落地运营","PC 后台"],
 ["村干部","移动履职（巡查/任务/审核）","干部移动端"],
 ["村民","掌上办事/参与治理","村民移动端"],
 ["参观/领导","大屏查看态势","数据大屏"]], [Cm(3), Cm(7), Cm(4)])
H2("4.2 角色权限矩阵（要点）")
para(doc, "遵循“组织隔离 + 角色授权 + 数据分级”：县级可见全县，乡级可见本乡，村级仅本村；功能权限由 admin/user/role-permission 配置；敏感操作（公开、积分、认养）需审核留痕。")
insert_diagram(doc, "rbac_architecture.png", "图 4-1 角色权限与组织隔离架构图")
H2("4.3 用户使用场景")
para(doc, "村民：打开 App 看通知、办低保、拍问题、兑积分；干部：巡村打卡、接诉求、审积分、写日志；县管：看大屏、发任务、审入驻、查报表。")

# ---------- 5 业务流程 ----------
H1("5. 业务流程")
para(doc, "平台核心业务按“申请/上报 → 受理分派 → 处置办理 → 审核留痕 → 办结反馈 → 数据沉淀”的主线运行，各业务域根据场景复用该主线。")
insert_diagram(doc, "business_process.png", "图 5-0 核心业务流程总图")
flows = [
 ("5.1 乡村开通流程", "乡镇/村提交入驻申请 → 县级审核资料 → 开通账号与千村千面配置 → 正式运营。", "flow_05_01_village_activation.png"),
 ("5.2 党建引领流程", "党组织发布活动/学习 → 党员参与 → 日志记录 → 组织生活统计。", "flow_05_02_party_building.png"),
 ("5.3 阳光村务流程", "起草三务公开内容 → 村级审核 → 发布公示 → 村民查看监督。", "flow_05_03_sunshine_village.png"),
 ("5.4 基层治理流程", "网格事件/随手拍上报 → 分派网格员 → 处置 → 上报闭环 → 治理大屏呈现。", "flow_05_04_grassroots_governance.png"),
 ("5.5 积分管理流程", "村民参与治理获积分 → 干部审核 → 积分发放 → 商城兑换 → 订单核销。", "flow_05_05_points_management.png"),
 ("5.6 乡村旅游流程", "村级维护旅游资源 → 平台展示 → 游客浏览/预订 → 村民增收。", "flow_05_06_rural_tourism.png"),
 ("5.7 农业认养流程", "发布认养项目 → 消费者认养 → 种养管理 → 产出配送 → 溯源查看。", "flow_05_07_agricultural_adoption.png"),
 ("5.8 表单应用流程", "村级用设计器建表 → 发布采集 → 村民填报 → 数据归集。", "flow_05_08_form_application.png"),
 ("5.9 设备管理流程", "设备台账登记 → 状态监测 → 告警处置 → 巡检维护。", "flow_05_09_device_management.png"),
 ("5.10 数据管理流程", "多源数据接入 → 清洗治理 → 指标计算 → 大屏/报表展示。", "flow_05_10_data_management.png"),
 ("5.11 千村千面配置流程", "选模板 → 配布局/配色/模块 → 预览 → 发布生效。", "flow_05_11_thousand_villages.png"),
 ("5.12 村民办事流程", "查指南 → 在线申请 → 受理审核 → 办理 → 结果回执。", "flow_05_12_villager_service.png"),
]
for t, d, img in flows:
    H2(t); para(doc, d)
    insert_diagram(doc, img, "图 " + t.replace("5.", "5-") + "图")

# ---------- 6 功能需求 ----------
H1("6. 功能需求")
para(doc, "以下按“县级 / 乡级 / 村级 / 大屏 / 村民移动 / 村干部 / 通用”七个分组展开，每个功能包含：功能说明、功能截图（原型全图）、功能字段及描述（不遗漏任何字段）、是否有业务流转。", size=10)
assigned = set()
for sec_key in ["county","town","village","dashboard","villager","cadre","common"]:
    H2(SECTION_TITLES[sec_key])
    for name, rels, desc, has_flow, flow_text in CATALOG[sec_key]:
        H3(name)
        sublabel(doc, "功能说明")
        para(doc, desc, size=10)
        sublabel(doc, "功能截图")
        add_image_list(doc, rels)
        sublabel(doc, "功能字段及描述")
        # merge field tables across rels (show each page)
        for rel in rels:
            if len(rels) > 1:
                para(doc, "▸ 页面：%s" % (DATA_JSON.get(rel, {}).get("main_title") or rel), size=9, color=RGBColor(0x33,0x66,0x99))
            field_table(doc, rel)
            assigned.add(rel)
        sublabel(doc, "是否有业务流转")
        if has_flow:
            para(doc, "是。" + (flow_text if flow_text else ""), size=10)
        else:
            para(doc, "否（本功能以信息查询/展示/配置为主，无多人多环节流转）。", size=10)

# coverage check
all_files = set(DATA_JSON.keys())
missing = all_files - assigned
extra = assigned - all_files
if missing:
    print("WARNING_UNCOVERED:", sorted(missing))
if extra:
    print("WARNING_EXTRA:", sorted(extra))

# ---------- 7 数据需求 ----------
H1("7. 数据需求")
H2("7.1 核心实体关系")
para(doc, "核心实体：组织（县/乡/村）→ 用户（干部/村民）→ 事务/诉求/事件 → 积分/订单 → 设备/产业/资源。组织为数据隔离根，用户归属组织，业务数据挂接组织与用户。")
insert_diagram(doc, "er_diagram.png", "图 7-1 核心数据实体关系图（ER）")
H2("7.2 核心数据实体清单")
make_table(doc, [["实体","关键属性（示例）","来源端"],
 ["组织","层级、名称、隶属、状态","admin/user/org-tree"],
 ["村民/农户","姓名、证件、家庭、承包地、标签","admin/user、village/my-village"],
 ["干部","姓名、职务、组织、联系方式","admin/user/cadre-list"],
 ["事务/诉求","类型、状态、分派、办结","admin/affairs、appeal、villager/report"],
 ["积分","余额、明细、规则","points 系列"],
 ["订单","商品、数量、状态、核销","points-shop"],
 ["设备","类型、位置、状态、告警","device 系列"],
 ["产业/资源","类别、规模、产值、地块","village/data、town/data"]], [Cm(3), Cm(8), Cm(3)])
H2("7.3 数据流")
para(doc, "采集（移动端/设备/IoT）→ 接入（datascreen 数据源）→ 治理（清洗/字典）→ 计算（指标）→ 服务（报表/大屏/移动端）。")
insert_diagram(doc, "dataflow_architecture.png", "图 7-3 数据架构与流向图")
H2("7.4 数据保留与归档")
para(doc, "业务流水保留 ≥ 3 年；日志审计保留 ≥ 6 个月；归档数据冷存储并可回溯。")
H2("7.5 数据质量要求")
para(doc, "完整性（必填校验）、准确性（字典约束）、一致性（组织隔离）、时效性（近实时大屏）。")

# ---------- 8 配置要求 ----------
H1("8. 配置要求")
for t, d in [
 ("8.1 开通配置", "县域开通、组织树初始化、账号与角色初始化（admin/city、user/org-tree）。"),
 ("8.2 业务配置", "便民事项、议事分类、报表模板、活动类型等业务参数。"),
 ("8.3 页面配置（千村千面）", "首页布局、配色、模块组合的村级定制（village/designer）。"),
 ("8.4 规则配置", "积分规则、消息模板、审核流开关（system/config）。"),
 ("8.5 设备配置", "IoT 设备接入、点位、告警阈值（device-manage）。"),
 ("8.6 角色与权限配置", "菜单/数据权限、最小权限（user/role-permission）。"),
 ("8.7 系统参数配置", "全局参数、字典、日志策略（system/*）。"),
]:
    H2(t); para(doc, d)

# ---------- 9 非功能要求 ----------
H1("9. 非功能要求")
for t, d in [
 ("9.1 性能要求", "列表页首屏 ≤ 1.5s；大屏数据刷新 ≤ 5s；并发支撑县级万级用户；批量报表导出 ≤ 30s。"),
 ("9.2 可用性要求", "年可用率 ≥ 99.5%；关键服务多副本；移动端弱网可降级。"),
 ("9.3 安全要求", "HTTPS 全链路、RBAC 权限、敏感数据脱敏（如身份证掩码）、操作审计、防越权（组织隔离）。"),
 ("9.4 兼容性要求", "PC 后台兼容 Chrome/Edge 最新版；移动端兼容 iOS12+/Android8+；大屏兼容主流拼接屏分辨率。"),
 ("9.5 可扩展性要求", "低代码表单、千村千面、数据源可插拔，支撑功能横向扩展。"),
 ("9.6 可维护性要求", "前后端分离、组件化、配置化，关键逻辑有日志与监控。"),
 ("9.7 移动端适配", "村民/干部端按移动视口（375–430px）设计，支持手势与拍照。"),
]:
    H2(t); para(doc, d)

# ---------- 10 技术架构 ----------
H1("10. 技术架构")
H2("10.1 总体架构")
para(doc, "表现层（四端：Vue/小程序/H5 + ECharts 大屏）→ 业务服务层（微服务/模块化）→ 数据层（关系库 + 时序/空间库）→ 基础设施（容器、网关、消息）。")
insert_diagram(doc, "tech_overall.png", "图 10-1 系统总体技术架构图")
H2("10.2 前端架构")
para(doc, "PC 后台与移动端 H5 采用组件化框架；大屏采用 ECharts + 大屏适配方案；共享组件库（shared/components-demo）保证一致性。")
insert_diagram(doc, "frontend_architecture.png", "图 10-2 前端架构图")
H2("10.3 后端架构")
para(doc, "按域划分服务：组织/用户、村务、治理、产业、积分、设备、数据；统一鉴权与网关。")
insert_diagram(doc, "backend_architecture.png", "图 10-3 后端微服务架构图")
H2("10.4 数据库设计")
para(doc, "关系库存业务主数据；空间库存地块/网格 GIS；时序库存设备监测；字典与配置独立表。")
H2("10.5 部署架构")
para(doc, "县级中心化部署 + 县/乡/村分级访问；大屏独立展示节点。")
insert_diagram(doc, "deploy_architecture.png", "图 10-5 部署架构图")
H2("10.6 集成架构")
para(doc, "通过数据源配置（datascreen/data-source）对接 IoT、第三方与上级平台；API 网关统一出口。")
insert_diagram(doc, "integration_architecture.png", "图 10-6 系统集成架构图")
H2("10.7 关键技术决策")
para(doc, "采用成熟开源栈、前后端分离、配置驱动与低代码，降低定制化成本，支撑“一县一档、千村千面”。")

# ---------- 11 运维部署 ----------
H1("11. 运维部署")
for t, d in [
 ("11.1 部署架构", "中心化私有云/政务云部署，反向代理 + 应用集群 + 数据库主从。"),
 ("11.2 部署清单", "应用服务器、数据库、缓存、对象存储、大屏展示机、域名/证书。"),
 ("11.3 部署流程", "环境准备 → 镜像/包发布 → 配置注入 → 健康检查 → 灰度 → 全量。"),
 ("11.4 容器化部署", "各服务容器化，编排管理，便于弹性伸缩。"),
 ("11.5 监控告警", "接口、资源、大屏数据延迟、设备离线等指标监控与分级告警。"),
 ("11.6 备份与恢复", "每日库备份 + 周全量；演练恢复 RTO<4h、RPO<1h。"),
 ("11.7 升级与回滚", "蓝绿/灰度发布，失败可快速回滚至上一稳定版本。"),
 ("11.8 安全运维", "漏洞扫描、日志审计、权限复核、等保合规。"),
]:
    H2(t); para(doc, d)

# ---------- 12 商业模式 ----------
H1("12. 商业模式")
H2("12.1 价值主张")
para(doc, "对政府：治理增效、考核可视、数据资产沉淀；对村民：办事不出村、参与有激励；对产业：文旅/认养引流增收。")
H2("12.2 收费与运营模式")
para(doc, "以政府统建采购（SaaS 订阅/项目制）为主；可选增值：认养农业交易佣金、文旅导流分成、积分商城供应链、数据增值服务（脱敏聚合）。")
H2("12.3 生态合作")
para(doc, "联合运营商（网络/物联网）、农商文旅商户、金融机构（惠民信贷）、硬件厂商共建乡村数字生态。")
H2("12.4 成本结构")
para(doc, "平台建设与运维、IoT 设备、运营推广、内容运营（三务/积分商城）为主要成本项。")

# ---------- 13 实施说明 ----------
H1("13. 实施说明")
H2("13.1 实施阶段")
make_table(doc, [["阶段","周期(建议)","交付物"],
 ["试点村建设","1–2 月","入驻+千村千面+首批功能上线"],
 ["县域推广","3–6 月","全县乡/村覆盖、大屏部署"],
 ["运营深化","持续","积分运营、产业对接、数据应用"]], [Cm(4), Cm(4), Cm(6)])
H2("13.2 组织保障")
para(doc, "成立县数字乡村工作专班，明确县-乡-村三级运营责任人；建立内容审核与数据安全责任制。")
H2("13.3 数据迁移与初始化")
para(doc, "组织树、村民/干部基础档案由存量数据导入 + 逐级核实；字典与规则统一初始化。")
H2("13.4 培训与推广")
para(doc, "对干部开展移动履职培训，对村民开展掌上办事引导，结合积分激励提升活跃。")
H2("13.5 风险与应对")
para(doc, "数据质量风险（强化审核）、活跃度风险（积分激励+内容运营）、安全风险（等保+审计）、依从性风险（合规审查）。")
H2("13.6 验收标准")
para(doc, "四端功能按本说明书第 6 章逐项验收；性能指标达第 9 章要求；安全通过等保测评。")

# ---------- 14 附录 ----------
H1("14. 附录")
H2("14.1 术语表")
terms = [
 ("数字乡村","以数字化驱动乡村产业、治理、服务现代化的综合体。"),
 ("三务公开","党务、村务、财务公开，落实阳光村务。"),
 ("千村千面","按村定制首页布局与风格的配置能力。"),
 ("积分制","以积分量化村民参与治理行为并兑换激励的机制。"),
 ("微网格","将村划分为最小治理单元，配网格员负责事件。"),
 ("农业认养","消费者线上认养农田/果树，远程种养并获产出。"),
 ("随手拍","村民拍照上报问题的轻量治理入口。"),
 ("数据大屏","面向指挥调度的大屏可视化展示。"),
 ("组织树","县-乡-村三级组织隶属结构，是权限与数据隔离基础。"),
 ("RBAC","基于角色的访问控制。"),
]
make_table(doc, [["术语","释义"]] + [[a,b] for a,b in terms], [Cm(4), Cm(10)])
H2("14.2 功能截图索引")
idx_rows = [["章节","功能","对应原型页面","截图文件"]]
for sec_key in ["county","town","village","dashboard","villager","cadre","common"]:
    for name, rels, desc, hf, ft in CATALOG[sec_key]:
        for rel in rels:
            shot = rel.replace("html/","").replace("/","_").replace(".html",".png")
            idx_rows.append([SECTION_TITLES[sec_key].split(" ",1)[-1], name, rel, shot])
make_table(doc, idx_rows, [Cm(2.5), Cm(4.5), Cm(4.5), Cm(3.0)])
H2("14.3 现有原型清单")
para(doc, "本说明书基于 html/ 目录下共 %d 个原型页面（admin %d、dashboard %d、villager %d、cadre %d、shared %d）编写。"
       % (len(DATA_JSON),
          sum(1 for k in DATA_JSON if k.startswith("admin")),
          sum(1 for k in DATA_JSON if k.startswith("dashboard")),
          sum(1 for k in DATA_JSON if k.startswith("villager")),
          sum(1 for k in DATA_JSON if k.startswith("cadre")),
          sum(1 for k in DATA_JSON if k.startswith("shared"))))
H2("14.4 修订记录")
make_table(doc, [["版本","日期","说明"],
 ["V1.0","2026-08-12","基于四端原型首版需求规格说明书（含截图与字段）"]], [Cm(2), Cm(3), Cm(9)])

# ---------- save ----------
doc.save(DOCX)
print("DOCX_SAVED", DOCX, "functions_total=", sum(len(v) for v in CATALOG.values()))
print("COVERAGE missing=", len(missing) if missing else 0, "extra=", len(extra) if extra else 0)
